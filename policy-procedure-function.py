#!/usr/bin/env python3
"""
Extract text from .eml, .pdf, .docx, .doc, and zipped folders into CSV.

Usage:
    python extract_files_to_csv.py /path/to/input_dir output.csv
"""

import csv
import email
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime
from email import policy
from email.parser import BytesParser
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".eml", ".pdf", ".docx", ".doc"}


def clean_text(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip()


def find_date_in_text(text: str) -> str:
    patterns = [
        r"\b\d{4}-\d{2}-\d{2}\b",
        r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
        r"\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b",
        r"\b[A-Za-z]{3,9}\s+\d{1,2},\s+\d{4}\b",
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)

    return ""


def find_created_by_in_text(text: str) -> str:
    patterns = [
        r"created by[:\s]+([^\n\r,;]+)",
        r"author[:\s]+([^\n\r,;]+)",
        r"prepared by[:\s]+([^\n\r,;]+)",
        r"written by[:\s]+([^\n\r,;]+)",
        r"from[:\s]+([^\n\r,;]+)",
        r"sent by[:\s]+([^\n\r,;]+)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return clean_text(match.group(1))

    return ""


def extract_eml(path: Path) -> dict:
    with open(path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    sender = msg.get("From", "")
    sent_date = msg.get("Date", "")
    subject = msg.get("Subject", "")

    body_parts = []

    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            disposition = part.get_content_disposition()

            if disposition == "attachment":
                continue

            if content_type == "text/plain":
                body_parts.append(part.get_content())
            elif content_type == "text/html":
                html = part.get_content()
                body_parts.append(BeautifulSoup(html, "html.parser").get_text(" "))
    else:
        content = msg.get_content()
        if msg.get_content_type() == "text/html":
            content = BeautifulSoup(content, "html.parser").get_text(" ")
        body_parts.append(content)

    text = clean_text(" ".join(body_parts))

    return {
        "text": text,
        "date": sent_date or find_date_in_text(text),
        "creator_or_sender": sender or find_created_by_in_text(text),
        "subject": subject,
    }


def extract_pdf(path: Path) -> dict:
    reader = PdfReader(str(path))
    pages = []

    for page in reader.pages:
        pages.append(page.extract_text() or "")

    text = clean_text(" ".join(pages))

    return {
        "text": text,
        "date": find_date_in_text(text),
        "creator_or_sender": find_created_by_in_text(text),
        "subject": "",
    }


def extract_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    text = clean_text(" ".join(paragraphs))

    return {
        "text": text,
        "date": find_date_in_text(text),
        "creator_or_sender": find_created_by_in_text(text),
        "subject": "",
    }


def extract_doc(path: Path) -> dict:
    """
    Legacy .doc support depends on antiword being installed.
    On Ubuntu/Debian:
        sudo apt install antiword
    """
    if not shutil.which("antiword"):
        return {
            "text": "",
            "date": "",
            "creator_or_sender": "",
            "subject": "",
            "error": "antiword not installed; cannot read legacy .doc file",
        }

    result = subprocess.run(
        ["antiword", str(path)],
        capture_output=True,
        text=True,
        errors="replace",
    )

    text = clean_text(result.stdout)

    return {
        "text": text,
        "date": find_date_in_text(text),
        "creator_or_sender": find_created_by_in_text(text),
        "subject": "",
        "error": result.stderr.strip(),
    }


def unzip_archives(input_dir: Path, work_dir: Path) -> Path:
    extracted_root = work_dir / "extracted"
    extracted_root.mkdir(parents=True, exist_ok=True)

    shutil.copytree(input_dir, extracted_root, dirs_exist_ok=True)

    processed_zips = set()

    while True:
        zip_files = [
            p for p in extracted_root.rglob("*.zip")
            if p not in processed_zips
        ]

        if not zip_files:
            break

        for zip_path in zip_files:
            target_dir = zip_path.with_suffix("")
            target_dir.mkdir(parents=True, exist_ok=True)

            try:
                with zipfile.ZipFile(zip_path, "r") as z:
                    z.extractall(target_dir)
                processed_zips.add(zip_path)
            except Exception:
                processed_zips.add(zip_path)

    return extracted_root


def process_file(path: Path) -> dict:
    ext = path.suffix.lower()

    try:
        if ext == ".eml":
            data = extract_eml(path)
        elif ext == ".pdf":
            data = extract_pdf(path)
        elif ext == ".docx":
            data = extract_docx(path)
        elif ext == ".doc":
            data = extract_doc(path)
        else:
            return {}

        return {
            "file_name": path.name,
            "file_path": str(path),
            "file_type": ext,
            "date_found_or_email_sent": data.get("date", ""),
            "creator_or_sender_found_in_file": data.get("creator_or_sender", ""),
            "email_subject": data.get("subject", ""),
            "text": data.get("text", ""),
            "error": data.get("error", ""),
        }

    except Exception as e:
        return {
            "file_name": path.name,
            "file_path": str(path),
            "file_type": ext,
            "date_found_or_email_sent": "",
            "creator_or_sender_found_in_file": "",
            "email_subject": "",
            "text": "",
            "error": str(e),
        }


def main():
    if len(sys.argv) != 3:
        print("Usage: python extract_files_to_csv.py /path/to/input_dir output.csv")
        sys.exit(1)

    input_dir = Path(sys.argv[1]).expanduser().resolve()
    output_csv = Path(sys.argv[2]).expanduser().resolve()

    if not input_dir.exists() or not input_dir.is_dir():
        raise ValueError(f"Input directory does not exist: {input_dir}")

    rows = []

    with tempfile.TemporaryDirectory() as temp_dir:
        work_dir = Path(temp_dir)
        extracted_root = unzip_archives(input_dir, work_dir)

        for path in extracted_root.rglob("*"):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                rows.append(process_file(path))

    fieldnames = [
        "file_name",
        "file_path",
        "file_type",
        "date_found_or_email_sent",
        "creator_or_sender_found_in_file",
        "email_subject",
        "text",
        "error",
    ]

    with open(output_csv, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    print(f"Done. Extracted {len(rows)} files to {output_csv}")


if __name__ == "__main__":
    main()
